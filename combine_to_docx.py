# combine the summary and detailed news into a docx file
# %%
from docx import Document
import shutil
import os
import datetime
from funcs import save_transcript_to_oss
todaydate=datetime.datetime.now().strftime('%Y-%m-%d %H:%M')

def combine_to_docx(project):
    print(project)
    doc = Document('format_new.docx')
    # Check if there's a matching summary md file
    summary_files = [f for f in sorted(os.listdir('./3_memo/'), reverse=True) if f.startswith(project)]
    if len(summary_files)>0:
        summary_md = open(f'./3_memo/{summary_files[0]}', 'r', encoding='utf-8').read()
        lines = summary_md.strip().split('\n')
        for line in lines:
            if line.startswith('# '):
                doc.add_paragraph(line.replace('# ',''), style='title1')
            elif line.startswith('## '):
                doc.add_paragraph(line.replace('## ',''), style='title2')
            elif line.startswith('### '):
                doc.add_paragraph(line.replace('### ',''), style='contentlist')
            elif len(line) > 10:
                doc.add_paragraph(line, style='sublist')

        doc.add_page_break()
    doc.add_paragraph('Full Discussion',style='section')


    raw_md=open(next((f'./2_wordforword/{f}' for f in sorted(os.listdir('./2_wordforword/'), reverse=True) if f.startswith(project)), f'./2_wordforword/{project}.txt'), 'r', encoding='utf-8').read()

    lines = raw_md.strip().split('\n')
    for line in lines:
        if line.startswith('# '):
            doc.add_paragraph(line.replace('# ',''), style='title1')
        elif line.startswith('## '):
            doc.add_paragraph(line.replace('## ',''), style='title2')
        elif line.startswith('%'):
            doc.add_paragraph(line.replace('%',''), style='boldcontent')
        elif line.startswith('$'):
            doc.add_paragraph(line.replace('$',''), style='nameformat')
        elif line.startswith('&'):
            doc.add_paragraph(line.replace('&',''), style='normaltranscript')
        elif len(line) > 20:
            doc.add_paragraph(line.replace('### ',''), style='contentlist')

    filename=f'{project} {todaydate}'
    doc.save(f'./4_docx/{filename}.docx')
    dest_path=f'~/Dropbox/VoiceMemos/{filename}.docx'
    dest_path=os.path.expanduser(dest_path)
    shutil.copyfile(f'./4_docx/{filename}.docx', dest_path)

    with open(f'./5_markdown/{filename}.md', 'w') as file:
        if len(summary_files)>0:
            file.write(f'# Key takeaways\n')
            lines = summary_md.strip().split('\n')
            for line in lines:
                if line.startswith('# '):
                    file.write(line.replace('# ','## ')+'\n')
                elif line.startswith('## '):
                    file.write(line.replace('## ','### ')+'\n')
                elif line.startswith('### '):
                    file.write(line.replace('### ','- ')+'\n')
                else:
                    file.write(line+'\n')
        file.write(f'\n# Full Discussion\n')
        for line in raw_md.strip().split('\n'):
            if line.startswith('%'):
                file.write('## '+line[1:]+'\n')
            elif line.startswith('$'):
                file.write('### '+line[1:]+'\n')
            elif line.startswith('&'):
                file.write(line[1:]+'\n')
            else:
                file.write(line+'\n')
        save_transcript_to_oss(os.path.expanduser(f'./5_markdown/{filename}.md'),f'{project}.md')
        
if __name__ == '__main__':
    project='weride 1q25 pc'
    combine_to_docx(project)